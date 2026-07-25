"""
Phase 0 - Document Digitization
Converts the raw source documents (docx, pdf, xlsx) supplied by the business
into Markdown files under docs/markdown/. This is a one-time / re-runnable
ingestion step. It does NOT interpret or invent rules - it only transcribes
structure (headings, paragraphs, tables) as faithfully as possible so that
Phase 1 (Knowledge Engineering) can work from Markdown instead of binary
formats.

Usage:
    .venv\\Scripts\\python.exe services\\digitize_documents.py
"""
from __future__ import annotations

import pathlib
import sys

import docx
import openpyxl
import pdfplumber

ROOT = pathlib.Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT.parent  # d:\Payroll_Validations
OUT_DIR = ROOT / "docs" / "markdown"

SOURCES = {
    "docx": [
        "AI_Payroll_Validation_Agent_Requirements.docx",
    ],
    "xlsx": [
        "MY Labour law and statutory calculation.xlsx",
    ],
    "pdf": [
        "EPF employee and employer contribution 10. Effective 1 October 2025.pdf",
        "SOCSO employee and employer NewContributionRateIncludingSKBBK.pdf",
        "EIS employee and employer 151124-Rate Contribution ACT 800.PDF",
    ],
}


def slugify(name: str) -> str:
    stem = pathlib.Path(name).stem
    stem = stem.replace(".", " ")
    return "-".join(stem.split()).lower()


def convert_docx(path: pathlib.Path) -> str:
    document = docx.Document(str(path))
    lines: list[str] = [f"# {path.stem}", "", f"> Source: `{path.name}` (converted verbatim, Phase 0 digitization)", ""]

    def table_to_md(table) -> list[str]:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            return []
        out = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * len(rows[0])) + " |"]
        for r in rows[1:]:
            out.append("| " + " | ".join(r) + " |")
        out.append("")
        return out

    # Walk the document body in order (paragraphs and tables interleaved)
    body = document.element.body
    table_iter = iter(document.tables)
    para_iter = iter(document.paragraphs)
    para_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}

    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            p = para_map.get(child)
            if p is None:
                continue
            text = p.text.strip()
            if not text:
                continue
            style = (p.style.name or "").lower() if p.style else ""
            if style.startswith("heading 1"):
                lines.append(f"## {text}")
            elif style.startswith("heading 2"):
                lines.append(f"### {text}")
            elif style.startswith("heading 3"):
                lines.append(f"#### {text}")
            elif style.startswith("list"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
            lines.append("")
        elif child.tag.endswith("}tbl"):
            t = table_map.get(child)
            if t is not None:
                lines.extend(table_to_md(t))

    # Images -> placeholders
    image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    if image_count:
        lines.append(f"> [IMAGE PLACEHOLDER x{image_count}] - images from the source document were not transcribed. Requires SME Validation if content is significant.")
        lines.append("")

    return "\n".join(lines)


def convert_xlsx(path: pathlib.Path) -> str:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    lines: list[str] = [f"# {path.stem}", "", f"> Source: `{path.name}` (converted verbatim, Phase 0 digitization)", ""]
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"## Sheet: {sheet_name}")
        lines.append("")
        rows = list(ws.iter_rows(values_only=True))
        rows = [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]
        if not rows:
            lines.append("_(empty sheet)_")
            lines.append("")
            continue
        max_cols = max(len(r) for r in rows)

        def cell(v):
            return "" if v is None else str(v).replace("\n", "<br>").strip()

        header = rows[0]
        header = list(header) + [None] * (max_cols - len(header))
        lines.append("| " + " | ".join(cell(c) or f"Col{i+1}" for i, c in enumerate(header)) + " |")
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        for r in rows[1:]:
            r = list(r) + [None] * (max_cols - len(r))
            lines.append("| " + " | ".join(cell(c) for c in r) + " |")
        lines.append("")
    return "\n".join(lines)


def convert_pdf(path: pathlib.Path, assets_dir: pathlib.Path) -> str:
    lines: list[str] = [f"# {path.stem}", "", f"> Source: `{path.name}` (converted verbatim, Phase 0 digitization)", ""]
    with pdfplumber.open(str(path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            lines.append(f"## Page {page_no}")
            lines.append("")
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text.strip())
                lines.append("")
            tables = page.extract_tables()
            found_table = False
            for t_idx, table in enumerate(tables, start=1):
                if not table:
                    continue
                found_table = True
                lines.append(f"**Table {page_no}.{t_idx}**")
                lines.append("")
                cleaned = [[("" if c is None else str(c).replace("\n", " ").strip()) for c in row] for row in table]
                if cleaned:
                    max_cols = max(len(r) for r in cleaned)
                    cleaned = [r + [""] * (max_cols - len(r)) for r in cleaned]
                    lines.append("| " + " | ".join(cleaned[0]) + " |")
                    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
                    for r in cleaned[1:]:
                        lines.append("| " + " | ".join(r) + " |")
                    lines.append("")
            images = getattr(page, "images", [])
            # If the page has no extractable text AND no extractable table, but does have raster
            # image content, it is very likely a scanned page. Render it to PNG so an SME (or a
            # human-in-the-loop OCR pass) can transcribe it later. Never invent the table content.
            if not text.strip() and not found_table and images:
                assets_dir.mkdir(parents=True, exist_ok=True)
                asset_name = f"{slugify(path.name)}-page{page_no}.png"
                try:
                    page.to_image(resolution=200).save(str(assets_dir / asset_name))
                    lines.append(
                        f"> **⚠ Requires SME Validation.** Page {page_no} appears to be a scanned "
                        f"image with no extractable text/table (found {len(images)} raster image(s)). "
                        f"No rate/rule figures were invented. Rendered for manual transcription at "
                        f"`docs/markdown/assets/{asset_name}`."
                    )
                except Exception as exc:  # pragma: no cover - best effort rendering
                    lines.append(
                        f"> **⚠ Requires SME Validation.** Page {page_no} appears to be a scanned "
                        f"image ({len(images)} raster image(s)) and could not be rendered automatically "
                        f"({exc}). Open the source PDF directly to transcribe."
                    )
                lines.append("")
            elif images and not found_table:
                lines.append(f"> [IMAGE PLACEHOLDER x{len(images)}] on page {page_no} - decorative/logo image, not transcribed.")
                lines.append("")
    return "\n".join(lines)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    converted = 0
    missing = []

    for fname in SOURCES["docx"]:
        src = SOURCE_DIR / fname
        if not src.exists():
            missing.append(fname)
            continue
        md = convert_docx(src)
        (OUT_DIR / f"{slugify(fname)}.md").write_text(md, encoding="utf-8")
        converted += 1
        print(f"[docx] {fname} -> {slugify(fname)}.md")

    for fname in SOURCES["xlsx"]:
        src = SOURCE_DIR / fname
        if not src.exists():
            missing.append(fname)
            continue
        md = convert_xlsx(src)
        out_name = f"{slugify(fname)}-xlsx.md"
        (OUT_DIR / out_name).write_text(md, encoding="utf-8")
        converted += 1
        print(f"[xlsx] {fname} -> {out_name}")

    for fname in SOURCES["pdf"]:
        src = SOURCE_DIR / fname
        if not src.exists():
            # try case-insensitive match
            candidates = list(SOURCE_DIR.glob("*.pdf")) + list(SOURCE_DIR.glob("*.PDF"))
            match = next((c for c in candidates if c.name.lower() == fname.lower()), None)
            if match is None:
                missing.append(fname)
                continue
            src = match
        md = convert_pdf(src, OUT_DIR / "assets")
        (OUT_DIR / f"{slugify(fname)}.md").write_text(md, encoding="utf-8")
        converted += 1
        print(f"[pdf]  {fname} -> {slugify(fname)}.md")

    # Also copy the already-plaintext .txt source verbatim for completeness
    txt_src = SOURCE_DIR / "MY Labour law and statutory calculation.txt"
    if txt_src.exists():
        text = txt_src.read_text(encoding="utf-8", errors="replace")
        md = f"# MY Labour law and statutory calculation (notes)\n\n> Source: `{txt_src.name}` (converted verbatim, Phase 0 digitization)\n\n```text\n{text}\n```\n"
        (OUT_DIR / "my-labour-law-and-statutory-calculation-txt.md").write_text(md, encoding="utf-8")
        converted += 1
        print(f"[txt]  {txt_src.name} -> my-labour-law-and-statutory-calculation-txt.md")

    print(f"\nConverted {converted} file(s) into {OUT_DIR}")
    if missing:
        print("MISSING (not found, skipped):")
        for m in missing:
            print(f"  - {m}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
